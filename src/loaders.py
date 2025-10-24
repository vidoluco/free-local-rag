"""
Document Loaders

Multi-format document loading for RAG system.
Supports: TXT, PDF, DOCX, Markdown, and web scraping.
"""

from pathlib import Path
from typing import List, Dict, Union
import re

# Document format libraries
import PyPDF2
import docx
import markdown
from bs4 import BeautifulSoup
import requests


class DocumentLoader:
    """Base class for document loading."""

    @staticmethod
    def load_txt(file_path: Union[str, Path]) -> str:
        """Load plain text file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    @staticmethod
    def load_pdf(file_path: Union[str, Path]) -> str:
        """Load PDF and extract text."""
        text_content = []

        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)

            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    text_content.append(f"=== PAGE {page_num} ===\n{text.strip()}")

        return "\n\n".join(text_content)

    @staticmethod
    def load_docx(file_path: Union[str, Path]) -> str:
        """Load DOCX and extract text."""
        doc = docx.Document(file_path)
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract table content
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    paragraphs.append(row_text)

        return "\n\n".join(paragraphs)

    @staticmethod
    def load_markdown(file_path: Union[str, Path]) -> str:
        """Load Markdown file and convert to plain text."""
        with open(file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # Convert markdown to HTML then extract text
        html = markdown.markdown(md_content)
        soup = BeautifulSoup(html, 'html.parser')

        # Get text while preserving some structure
        text = soup.get_text(separator='\n')

        # Clean up extra whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)

        return text.strip()

    @staticmethod
    def scrape_webpage(url: str, selector: str = None) -> str:
        """
        Scrape content from webpage.

        Args:
            url: URL to scrape
            selector: Optional CSS selector to extract specific content

        Returns:
            Extracted text content
        """
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }

        response = requests.get(url, headers=headers, timeout=30)
        response.raise_for_status()

        soup = BeautifulSoup(response.content, 'html.parser')

        # Remove script and style elements
        for script in soup(["script", "style", "nav", "footer"]):
            script.decompose()

        # Extract content based on selector or get all text
        if selector:
            elements = soup.select(selector)
            text = "\n\n".join(el.get_text(strip=True) for el in elements)
        else:
            text = soup.get_text(separator='\n')

        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        text = '\n'.join(line for line in lines if line)

        return text

    @classmethod
    def load_document(cls, file_path: Union[str, Path]) -> Dict[str, str]:
        """
        Load document and return content with metadata.

        Args:
            file_path: Path to document file

        Returns:
            Dict with 'content', 'filename', 'format' keys
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Determine format and load
        ext = file_path.suffix.lower()

        loaders = {
            '.txt': cls.load_txt,
            '.pdf': cls.load_pdf,
            '.docx': cls.load_docx,
            '.doc': cls.load_docx,
            '.md': cls.load_markdown,
            '.markdown': cls.load_markdown,
        }

        if ext not in loaders:
            raise ValueError(
                f"Unsupported file format: {ext}. "
                f"Supported: {', '.join(loaders.keys())}"
            )

        content = loaders[ext](file_path)

        return {
            'content': content,
            'filename': file_path.name,
            'format': ext.lstrip('.'),
            'path': str(file_path)
        }

    @classmethod
    def load_directory(cls, dir_path: Union[str, Path]) -> List[Dict[str, str]]:
        """
        Load all supported documents from directory.

        Args:
            dir_path: Path to directory containing documents

        Returns:
            List of document dicts with content and metadata
        """
        dir_path = Path(dir_path)

        if not dir_path.exists():
            raise FileNotFoundError(f"Directory not found: {dir_path}")

        supported_exts = {'.txt', '.pdf', '.docx', '.doc', '.md', '.markdown'}
        documents = []

        for file_path in dir_path.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in supported_exts:
                try:
                    doc = cls.load_document(file_path)
                    documents.append(doc)
                    print(f"✓ Loaded: {file_path.name}")
                except Exception as e:
                    print(f"✗ Error loading {file_path.name}: {e}")

        return documents


class WebScraper:
    """Specialized web scraping utilities."""

    @staticmethod
    def scrape_multiple_urls(urls: List[str], selector: str = None) -> List[Dict[str, str]]:
        """
        Scrape multiple URLs and return structured content.

        Args:
            urls: List of URLs to scrape
            selector: Optional CSS selector

        Returns:
            List of dicts with 'content', 'url', 'title'
        """
        results = []

        for url in urls:
            try:
                print(f"Scraping: {url}")
                content = DocumentLoader.scrape_webpage(url, selector)

                # Try to extract title
                response = requests.get(url, timeout=30)
                soup = BeautifulSoup(response.content, 'html.parser')
                title = soup.title.string if soup.title else url

                results.append({
                    'content': content,
                    'url': url,
                    'title': title.strip(),
                    'format': 'web'
                })
                print(f"✓ Scraped: {title}")

            except Exception as e:
                print(f"✗ Error scraping {url}: {e}")

        return results

    @staticmethod
    def scrape_from_config(config_data: Dict) -> List[Dict[str, str]]:
        """
        Scrape URLs specified in config.yaml.

        Args:
            config_data: Configuration dict with web_scraping section

        Returns:
            List of scraped content dicts
        """
        web_config = config_data.get('data', {}).get('web_scraping', {})

        if not web_config.get('enabled', False):
            return []

        urls = web_config.get('urls', [])
        selector = web_config.get('selector')

        if not urls:
            return []

        return WebScraper.scrape_multiple_urls(urls, selector)


class ContentAggregator:
    """Aggregate content from multiple sources into unified format."""

    @staticmethod
    def aggregate_documents(documents: List[Dict[str, str]]) -> str:
        """
        Combine multiple documents into single text with section markers.

        Args:
            documents: List of document dicts from loaders

        Returns:
            Combined text with section markers for chunking
        """
        sections = []

        for doc in documents:
            filename = doc.get('filename', doc.get('title', doc.get('url', 'Unknown')))
            content = doc.get('content', '')

            if content.strip():
                section = f"=== DOCUMENT: {filename} ===\n\n{content.strip()}"
                sections.append(section)

        return "\n\n".join(sections)

    @staticmethod
    def save_aggregated_content(
        documents: List[Dict[str, str]],
        output_path: Union[str, Path]
    ):
        """
        Save aggregated content to file.

        Args:
            documents: List of document dicts
            output_path: Path to save combined content
        """
        combined_content = ContentAggregator.aggregate_documents(documents)

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(combined_content)

        print(f"✓ Aggregated content saved: {output_path}")
        print(f"  Total documents: {len(documents)}")
        print(f"  Total characters: {len(combined_content)}")
